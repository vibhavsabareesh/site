#!/usr/bin/env python3
"""Turn a supplied background guide (PDF or DOCX) into content.json.

Headings are detected by font size (anything meaningfully larger than the
document's body size) or by a short, fully bold line. Inline bold inside a
paragraph is preserved as [text, bold] runs so lead-ins survive.

Usage:  python extract.py in.pdf|in.docx out.json --committee UNSC --agenda "..." [--emblem emblems/unsc.png] [--skip 1,2]
"""
import argparse, collections, json, os, re, sys

# Word and Google Docs export bullets as private-use glyphs from symbol fonts
# (\uf0b7 and friends); those have no Montserrat glyph and would render as tofu.
BULLET_CHARS = "•●▪‣◦○➢➤·∙"
PUA_BULLETS = "\uf0a7\uf0b7\uf0d8\uf06c\uf0fc\uf0a8\uf076"
# a numeric marker is at most two digits: "12." is a list item, "5334." is
# a bill number that happens to start a line
BULLET = re.compile(r"^\s*(?:[" + BULLET_CHARS + PUA_BULLETS + r"\-\u2013]|\d{1,2}[.)])\s*")


def strip_marker(runs, n):
    """Drop the first n characters of a paragraph, spanning runs if needed."""
    out, left = [], n
    for text, bold in runs:
        if left <= 0:
            out.append([text, bold])
        elif len(text) <= left:
            left -= len(text)
        else:
            out.append([text[left:], bold])
            left = 0
    return out or [["", False]]


def merge(runs):
    out = []
    for text, bold in runs:
        if out and not text.strip():
            out[-1][0] += text          # keep spacing with the run before it
            continue
        if out and out[-1][1] == bold:
            out[-1][0] += text
        else:
            out.append([text, bold])
    return [[t, b] for t, b in out if t.strip() or len(out) == 1]


def downsample(path, placed_w_pt, dpi=200, quality=82):
    """Shrink an image to the resolution it is actually shown at.

    Photo-led guides embed originals many times larger than their placed size;
    left alone the rebuilt PDF runs to tens of megabytes.
    """
    try:
        from PIL import Image
    except ImportError:
        return path
    try:
        im = Image.open(path)
    except Exception:
        return path
    target = max(int(placed_w_pt * dpi / 72), 32)
    if im.width > target:
        im = im.resize((target, max(1, round(im.height * target / im.width))), Image.LANCZOS)
    has_alpha = im.mode in ("RGBA", "LA") or "transparency" in im.info
    if has_alpha:
        im.save(path)
        return path
    jpg = os.path.splitext(path)[0] + ".jpg"
    im.convert("RGB").save(jpg, quality=quality, optimize=True)
    if jpg != path and os.path.exists(path):
        os.remove(path)
    return jpg


def save_image(doc, xref, path):
    """Write an image out, composing its soft mask so transparency survives.

    Without this a masked image (a watermark, say) is drawn as an opaque
    rectangle over the page.
    """
    import pymupdf
    info = doc.extract_image(xref)
    if info.get("smask"):
        pix = pymupdf.Pixmap(info["image"])
        mask = pymupdf.Pixmap(doc.extract_image(info["smask"])["image"])
        try:
            pix = pymupdf.Pixmap(pix, mask)
        except Exception:
            pass
        path = os.path.splitext(path)[0] + ".png"
        pix.save(path)
        return path
    with open(path, "wb") as fh:
        fh.write(info["image"])
    return path


def detect_tables(lines, min_rows=3, gap=18.0):
    """Find borderless tables by geometry.

    Guides set tables as plain text in columns, so there are no rules to find.
    Each cell arrives as its own line sharing a baseline with its neighbours,
    so rows are rebuilt by y and a table is several consecutive rows whose
    cells start at the same x positions.

    Returns {first_line_index: table} and the set of line indices consumed.
    """
    # group lines into visual rows by baseline
    order, rows = [], []
    for i, l in enumerate(lines):
        if rows and l["page"] == lines[rows[-1][0]]["page"] \
                and abs(l["y"] - lines[rows[-1][0]]["y"]) <= 3:
            rows[-1].append(i)
        else:
            rows.append([i])
    for r in rows:
        r.sort(key=lambda i: lines[i]["x0"])

    counts = collections.Counter()
    for r in rows:
        for i in r:
            counts[round(lines[i]["x0"] / 4) * 4] += 1
    anchors = sorted(x for x, n in counts.items() if n >= min_rows)

    def anchor_of(x):
        best = min(anchors, key=lambda a: abs(a - x), default=None)
        return best if best is not None and abs(best - x) <= 6 else None

    tables, used, ri = {}, set(), 0
    while ri < len(rows):
        cols = [anchor_of(lines[i]["x0"]) for i in rows[ri]]
        cand = sorted({c for c in cols if c is not None})
        wide = len(cand) >= 2 and min(b - a for a, b in zip(cand, cand[1:])) >= gap
        if len(rows[ri]) < 2 or not wide:
            ri += 1
            continue
        page = lines[rows[ri][0]]["page"]
        gaps = [lines[rows[k + 1][0]]["y"] - lines[rows[k][0]]["y"]
                for k in range(ri, min(ri + 6, len(rows) - 1))
                if lines[rows[k + 1][0]]["page"] == page]
        gaps = [g for g in gaps if g > 0]
        typical = sorted(gaps)[len(gaps) // 2] if gaps else 20.0
        out, rj = [], ri
        while rj < len(rows) and lines[rows[rj][0]]["page"] == page:
            if rj > ri:
                step = lines[rows[rj][0]]["y"] - lines[rows[rj - 1][0]]["y"]
                if step > typical * 2.5:
                    break        # a gap this big means the table has ended
            if any(lines[i]["kind"][0] for i in rows[rj]):
                break            # a heading is not a table row
            cj = [anchor_of(lines[i]["x0"]) for i in rows[rj]]
            if all(c is None for c in cj):
                break
            starts = cj[0] == cand[0]
            if cj[0] is not None and cj[0] < cand[0]:
                break            # text to the left of the table: it has ended
            if starts:
                out.append({c: [] for c in cand})
            elif not out:
                break
            for i, c in zip(rows[rj], cj):
                out[-1].setdefault(c if c is not None else cand[-1], []).append(i)
            rj += 1
        if len(out) >= min_rows:
            first = min(rows[ri])
            # the header row often spans fewer columns than the body, so the
            # column list is the union across every row
            allc = sorted({c for row in out for c in row
                           if any(r.get(c) for r in out)})
            tables[first] = {"anchors": allc, "rows": out}
            for r in rows[ri:rj]:
                used.update(r)
            ri = rj
        else:
            ri += 1
    return tables, used


def layout(doc, skip, outdir):
    """Capture a page as geometry: every image and text span with its position.

    Used for guides whose layout *is* the content -- photo grids and the like --
    where reflowing the text would destroy the page.
    """
    import os
    os.makedirs(outdir, exist_ok=True)
    pages = []
    for pno, pg in enumerate(doc, start=1):
        if pno in skip:
            continue
        items = []
        for n, inf in enumerate(pg.get_image_info(xrefs=True)):
            x0, y0, x1, y1 = inf["bbox"]
            if x1 - x0 < 24 or y1 - y0 < 24:
                continue
            img = doc.extract_image(inf["xref"])
            name = f"p{pno}_{n}.{img['ext']}"
            saved = save_image(doc, inf["xref"], os.path.join(outdir, name))
            saved = downsample(saved, x1 - x0)
            items.append({"k": "img", "src": saved, "rect": [x0, y0, x1, y1]})
        for b in pg.get_text("dict")["blocks"]:
            for l in b.get("lines", []):
                for s in l["spans"]:
                    if not s["text"].strip():
                        continue
                    items.append({"k": "txt", "text": s["text"],
                                  "x": s["origin"][0], "y": s["origin"][1],
                                  "size": s["size"], "bold": "Bold" in s["font"],
                                  "w": s["bbox"][2] - s["bbox"][0]})
        if items:
            pages.append({"page": pno, "items": items})
    return pages


def figures(doc, skip, outdir):
    """Save the source's own images so they can be placed back in the rebuild."""
    import os
    out = []
    if outdir:
        os.makedirs(outdir, exist_ok=True)
    for pno, pg in enumerate(doc, start=1):
        if pno in skip:
            continue
        for n, inf in enumerate(pg.get_image_info(xrefs=True)):
            x0, y0, x1, y1 = inf["bbox"]
            if x1 - x0 < 40 or y1 - y0 < 40:
                continue                       # rules, bullets, artefacts
            img = doc.extract_image(inf["xref"])
            name = f"p{pno}_{n}.{img['ext']}"
            path = downsample(save_image(doc, inf["xref"], os.path.join(outdir, name)),
                              x1 - x0)
            out.append({"t": "i", "src": path, "page": pno, "y": y0,
                        "w": x1 - x0, "h": y1 - y0, "runs": [["", False]]})
    return out


def from_pdf(path, skip):
    import pymupdf
    doc = pymupdf.open(path)
    sizes = collections.Counter()
    for pg in doc:
        for b in pg.get_text("dict")["blocks"]:
            for l in b.get("lines", []):
                for s in l["spans"]:
                    if s["text"].strip():
                        sizes[round(s["size"], 1)] += len(s["text"])
    body_size = sizes.most_common(1)[0][0]

    # collect every line with its geometry first; many guides put each line in
    # its own block, so paragraphs have to be reassembled from the text edges
    lines = []
    for pno, pg in enumerate(doc, start=1):
        if pno in skip:
            continue
        for b in pg.get_text("dict")["blocks"]:
            for l in b.get("lines", []):
                spans = [s for s in l["spans"] if s["text"].strip()]
                if not spans:
                    continue
                text = "".join(s["text"] for s in spans).strip()
                maxsize = max(round(s["size"], 1) for s in spans)
                runs = [[s["text"], "Bold" in s["font"]] for s in spans]
                major = maxsize >= body_size + 1.4
                kind = (True, 1) if major else (False, 1)
                segs = []
                for s in spans:
                    if segs and s["bbox"][0] - segs[-1]["x1"] < 8:
                        segs[-1]["runs"].append([s["text"], "Bold" in s["font"]])
                        segs[-1]["x1"] = s["bbox"][2]
                    else:
                        segs.append({"x0": s["bbox"][0], "x1": s["bbox"][2],
                                     "runs": [[s["text"], "Bold" in s["font"]]]})
                lines.append({"text": text, "runs": runs, "kind": kind, "segs": segs,
                              "x0": l["bbox"][0], "x1": l["bbox"][2],
                              "y": l["bbox"][1], "page": pno})
    if not lines:
        return []
    # Text set in a narrow column (beside a figure, say) has its own right
    # edge; judging those lines against the page's widest line would end a
    # paragraph on every line. So each line is measured against the column it
    # belongs to: the widest line nearby that starts at the same indent.
    # Normally a paragraph ends when its last line stops short of the page's
    # right edge. Text wrapped beside a figure has its own, much narrower edge,
    # so those lines are measured against their column instead.
    body = [l["x1"] for l in lines if not l["kind"][0]] or [l["x1"] for l in lines]
    page_right = max(body)
    figure_bands = {}
    for pno, pg in enumerate(doc, start=1):
        if pno in skip:
            continue
        bands = []
        for inf in pg.get_image_info():
            x0, y0, x1, y1 = inf["bbox"]
            if x1 - x0 >= 40 and y1 - y0 >= 40:
                bands.append((y0 - 4, y1 + 4))
        if bands:
            figure_bands[pno] = bands
    for l in lines:
        band = next((b for b in figure_bands.get(l["page"], [])
                     if b[0] <= l["y"] <= b[1]), None)
        if band is None:
            l["right"] = page_right
        else:
            beside = [o["x1"] for o in lines
                      if o["page"] == l["page"] and band[0] <= o["y"] <= band[1]
                      and abs(o["x0"] - l["x0"]) <= 8]
            l["right"] = max(beside or [l["x1"]])
    page_left = min(l["x0"] for l in lines)

    paras, cur, cur_kind = [], [], None
    cur_page, cur_y = [lines[0]["page"]], [lines[0]["y"]]

    def flush():
        if not cur:
            return
        runs = merge(cur)
        text = "".join(r[0] for r in runs).strip()
        if text.isdigit() and len(text) <= 3:
            cur.clear()                       # a bare page number
            return
        if text:
            major = cur_kind[0]
            all_bold = all(b for s, b in runs if s.strip())
            # a whole short paragraph in bold is a subheading; bold inside a
            # longer paragraph is just emphasis
            # a whole short paragraph set in bold is a subheading; guides
            # often set these only a point larger than the body
            sub = not major and all_bold and len(text) < 70 and len(text.split()) <= 10
            paras.append({"text": text, "runs": runs,
                          "heading": major or sub, "level": 1 if major else 2,
                          "page": cur_page[0], "y": cur_y[0]})
        cur.clear()

    tables, table_lines = detect_tables(lines)

    for i, l in enumerate(lines):
        if i in tables:
            flush()
            tb = tables[i]
            rows = []
            for row in tb["rows"]:
                cells = []
                for a in tb["anchors"]:
                    cell, prev = [], None
                    col_right = (tb["anchors"][tb["anchors"].index(a) + 1]
                                 if a != tb["anchors"][-1] else page_right)
                    for idx in row.get(a, []):
                        ln = lines[idx]
                        if cell:
                            # if the next line's first word would have fitted on
                            # the previous one, the break was mid-word
                            head = ln["text"].split(" ")[0]
                            room = col_right - prev["x1"]
                            pw = prev["x1"] - prev["x0"]
                            char_w = pw / max(len(prev["text"]), 1)
                            midword = len(head) <= 4 and room > len(head) * char_w * 0.8
                            if not midword:
                                cell.append([" ", False])
                        cell.extend(ln["runs"])
                        prev = ln
                    cells.append(merge(cell) if cell else [["", False]])
                rows.append(cells)
            head = all(len("".join(x[0] for x in cell).strip()) <= 30 for cell in rows[0]) \
                if rows else False
            paras.append({"t": "tbl", "anchors": tb["anchors"], "rows": rows,
                          "header": head,
                          "text": "", "runs": [["", False]], "heading": False,
                          "level": 1, "page": l["page"], "y": l["y"]})
            continue
        if i in table_lines:
            continue
        span = max(l["right"] - page_left, 1.0)
        right_edge = l["right"]
        if cur and (l["kind"] != cur_kind or cur_kind[0]):
            flush()
        if not cur:
            cur_page[0], cur_y[0] = l["page"], l["y"]
        cur_kind = l["kind"]
        cur.extend(l["runs"])
        cur.append([" ", l["runs"][-1][1]])
        nxt = lines[i + 1] if i + 1 < len(lines) else None
        short = right_edge - l["x1"]
        # ragged-right text: a break needs terminal punctuation and a gap, or a
        # very short line. Ending mid-sentence means the paragraph continues.
        ends_para = (
            l["kind"][0]
            or nxt is None
            or nxt["kind"] != l["kind"]
            or bool(BULLET.match(nxt["text"]))
            or (l["text"].rstrip().endswith((".", "!", "?", ":", ";", '"', "”")) and short > 0.10 * span)
            or short > 0.34 * span
        )
        if ends_para:
            flush()
    flush()
    return paras


DOCX_NS = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main",
           "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
           "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
EMU_PER_PT = 12700


def docx_images(para, doc, outdir, counter):
    """Inline images in this paragraph, saved and sized in points."""
    import os
    out = []
    for blip in para._p.findall(".//{%s}blip" % DOCX_NS["a"]):
        rid = blip.get("{%s}embed" % DOCX_NS["r"])
        if not rid or rid not in doc.part.related_parts:
            continue
        part = doc.part.related_parts[rid]
        ext = os.path.splitext(part.partname)[1] or ".png"
        counter[0] += 1
        name = f"docx_{counter[0]}{ext}"
        os.makedirs(outdir, exist_ok=True)
        path = os.path.join(outdir, name)
        with open(path, "wb") as fh:
            fh.write(part.blob)
        holder = blip.getparent().getparent().getparent()
        ext_el = holder.find(".//{%s}extent" % DOCX_NS["wp"]) if holder is not None else None
        w = h = 0
        if ext_el is not None:
            w = int(ext_el.get("cx", 0)) / EMU_PER_PT
            h = int(ext_el.get("cy", 0)) / EMU_PER_PT
        out.append({"t": "i", "src": path, "w": w, "h": h, "runs": [["", False]]})
    return out


def from_docx(path, skip, figdir=""):
    import docx
    d = docx.Document(path)
    counter = [0]
    sizes = collections.Counter()
    for p in d.paragraphs:
        for r in p.runs:
            if r.text.strip() and r.font.size:
                sizes[r.font.size.pt] += len(r.text)
    body_size = sizes.most_common(1)[0][0] if sizes else 11.0
    paras = []
    for p in d.paragraphs:
        if figdir:
            for img in docx_images(p, d, figdir, counter):
                paras.append(img)
        text = p.text.strip()
        if not text:
            continue
        runs = merge([[r.text, bool(r.bold)] for r in p.runs if r.text]) or [[text, False]]
        joined = "".join(r[0] for r in runs)
        if len(joined.split()) < len(text.split()):
            # hyperlink runs live outside paragraph.runs; keep the full text
            runs = [[text, all(b for _, b in runs)]]
        maxsize = max([r.font.size.pt for r in p.runs if r.font.size] or [body_size])
        all_bold = all(b for _, b in runs if _.strip())
        major = maxsize >= body_size + 1.5 and all_bold
        heading = major or (maxsize >= body_size + 1.0 and all_bold)
        paras.append({"text": text, "runs": runs, "heading": heading,
                      "level": 1 if major else 2})
    return paras


def build_pages(paras, committee, agenda, emblem):
    """One entry per source page, so the output can be a 1:1 replica."""
    pages, cur_no, cur = [], None, None
    for p in paras:
        if p.get("page") != cur_no:
            cur_no = p.get("page")
            cur = {"page": cur_no, "blocks": []}
            pages.append(cur)
        if p.get("t") == "tbl":
            cur["blocks"].append({"t": "tbl", "anchors": p["anchors"], "rows": p["rows"],
                                  "header": p.get("header", False)})
            continue
        if p.get("t") == "i":
            cur["blocks"].append({k: p[k] for k in ("t", "src", "w", "h", "runs")})
            continue
        if p["heading"]:
            cur["blocks"].append({"t": "h", "text": p["text"],
                                  "level": p.get("level", 1), "runs": p["runs"]})
            continue
        runs = p["runs"]
        m = BULLET.match(p["text"])
        if m:
            marker = m.group(0).strip()
            runs = strip_marker(runs, len(m.group(0)))
            blk = {"t": "b", "runs": runs}
            if marker and marker[0].isdigit():
                blk["marker"] = marker
            cur["blocks"].append(blk)
        else:
            cur["blocks"].append({"t": "p", "runs": runs})
    pages = [pg for pg in pages if any(
        b.get("t") in ("i", "tbl") or any(r[0].strip() for r in b["runs"])
        for b in pg["blocks"])]
    out = {"committee": committee, "agenda": agenda, "pages": pages}
    if emblem:
        out["emblem"] = emblem
    return out


def build(paras, committee, agenda, emblem):
    sections, cur = [], None

    def ensure():
        nonlocal cur
        if cur is None:
            cur = {"heading": "", "level": 1, "page_break": False, "blocks": []}
            sections.append(cur)
        return cur

    for p in paras:
        if p.get("t") == "tbl":
            ensure()["blocks"].append({"t": "tbl", "anchors": p["anchors"],
                                       "rows": p["rows"], "header": p.get("header", False)})
            continue
        if p.get("t") == "i":
            ensure()["blocks"].append({k: p[k] for k in ("t", "src", "w", "h", "runs")})
            continue
        if p["heading"]:
            cur = {"heading": p["text"], "level": p.get("level", 1),
                   "page_break": False, "blocks": []}
            sections.append(cur)
            continue
        ensure()
        runs = p["runs"]
        m = BULLET.match(p["text"])
        if m:
            marker = m.group(0).strip()
            runs = strip_marker(runs, len(m.group(0)))
            blk = {"t": "b", "runs": runs}
            if marker and marker[0].isdigit():
                blk["marker"] = marker      # numbered lists keep their numbers
            cur["blocks"].append(blk)
        else:
            cur["blocks"].append({"t": "p", "runs": runs})
    out = {"committee": committee, "agenda": agenda, "sections": sections}
    if emblem:
        out["emblem"] = emblem
    return out


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("src"); ap.add_argument("out")
    ap.add_argument("--committee", required=True)
    ap.add_argument("--agenda", required=True)
    ap.add_argument("--emblem")
    ap.add_argument("--skip", default="", help="1-based source pages to drop, e.g. cover/TOC")
    ap.add_argument("--start", default="", help="drop everything before the paragraph starting with this")
    ap.add_argument("--figures", default="", help="directory to save the source's images into")
    ap.add_argument("--layout", action="store_true",
                    help="reproduce the source page geometry (photo grids etc.) "
                         "instead of reflowing the text")
    ap.add_argument("--flow", action="store_true",
                    help="flow continuously instead of mirroring the source pagination")
    a = ap.parse_args()
    skip = {int(x) for x in a.skip.split(",") if x.strip()}
    if a.layout:
        import pymupdf
        doc = pymupdf.open(a.src)
        content = {"committee": a.committee, "agenda": a.agenda,
                   "layout_pages": layout(doc, skip, a.figures or "figures")}
        if a.emblem:
            content["emblem"] = a.emblem
        json.dump(content, open(a.out, "w", encoding="utf-8"), indent=1, ensure_ascii=False)
        items = [i for pg in content["layout_pages"] for i in pg["items"]]
        print(f"{a.out}: {len(content['layout_pages'])} pages laid out, "
              f"{sum(1 for i in items if i['k'] == 'img')} images, "
              f"{sum(1 for i in items if i['k'] == 'txt')} text spans")
        raise SystemExit

    is_docx = a.src.lower().endswith(".docx")
    paras = from_docx(a.src, skip, a.figures) if is_docx else from_pdf(a.src, skip)
    figs = []
    if not is_docx and a.figures:
        import pymupdf
        figs = figures(pymupdf.open(a.src), skip, a.figures)
        print(f"  {len(figs)} figure(s) carried over")
    if a.start:
        for i, p in enumerate(paras):
            if p["text"].lower().startswith(a.start.lower()):
                paras = paras[i:]
                break
    if figs:
        paras = sorted(paras + figs, key=lambda p: (p.get("page", 0), p.get("y", 0)))
    paged = any("page" in p for p in paras) and not a.flow
    content = (build_pages if paged else build)(paras, a.committee, a.agenda, a.emblem)
    json.dump(content, open(a.out, "w", encoding="utf-8"), indent=1, ensure_ascii=False)
    if "pages" in content:
        blocks = [b for pg in content["pages"] for b in pg["blocks"]]
        ntbl = sum(1 for b in blocks if b["t"] == "tbl")
        if ntbl:
            print(f"  {ntbl} table(s) detected")
        print(f"{a.out}: {len(content['pages'])} source pages, "
              f"{sum(1 for b in blocks if b['t'] == 'h')} headings, "
              f"{sum(1 for b in blocks if b['t'] == 'p')} paragraphs, "
              f"{sum(1 for b in blocks if b['t'] == 'b')} bullets")
    else:
        blocks = [b for s in content["sections"] for b in s["blocks"]]
        print(f"{a.out}: {len(content['sections'])} sections, "
              f"{sum(1 for b in blocks if b['t'] == 'p')} paragraphs, "
              f"{sum(1 for b in blocks if b['t'] == 'b')} bullets")
